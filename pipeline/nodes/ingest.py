import os
from typing import List
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from models.schemas import Document


def _read_pdf(path: str) -> tuple[str, int]:
    doc = fitz.open(path)
    pages = doc.page_count
    texts = []
    for i in range(pages):
        page = doc.load_page(i)
        texts.append(page.get_text("text"))
    return "\n".join(texts), pages


def _read_docx(path: str) -> tuple[str, int]:
    d = DocxDocument(path)
    parts = [p.text for p in d.paragraphs]
    return "\n".join(parts), 0


def _read_txt(path: str) -> tuple[str, int]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read(), 0


def load_files(file_paths: List[str]) -> List[Document]:
    documents: List[Document] = []
    for p in file_paths:
        ext = os.path.splitext(p)[1].lower()
        text = ""
        pages = 0
        if ext == ".pdf":
            text, pages = _read_pdf(p)
        elif ext == ".docx":
            text, pages = _read_docx(p)
        elif ext == ".txt":
            text, pages = _read_txt(p)
        else:
            continue
        documents.append(Document(name=os.path.basename(p), path=p, text=text, pages=pages))
    return documents
